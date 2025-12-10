import streamlit as st
import docx
import fitz
import re
import requests
import json
from io import BytesIO


# ===============================================================
# 1. Extract text from a single PDF
# ===============================================================
def extract_pdf_text(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for p in doc:
        text += p.get_text() + "\n"
    return text


# ===============================================================
# 2. Extract and combine multiple PDFs
# ===============================================================
def extract_multiple_pdfs(pdf_files):
    combined = ""
    for pdf in pdf_files:
        combined += f"\n\n--- PDF FILE: {pdf.name} ---\n"
        combined += extract_pdf_text(pdf.read())
    return combined


# ===============================================================
# 3. CALL LLM (DeepSeek via OpenRouter)
# ===============================================================
def call_llm(template_text, pdf_text):

    api_key = st.secrets["OPENROUTER_API_KEY"]

    SYSTEM_PROMPT = """
You are an insurance GLR automation assistant.

Your job:
1. Read the user-provided GLR template text (DOCX converted to raw text).
2. Read the combined PDF photo report text.
3. Detect ALL placeholders dynamically:
   - Bracket placeholders: [FIELD], [DATE_LOSS], [INSURED_NAME]
   - Parentheses placeholders: ( ... )
   - Quote placeholders: "..." or “...”
   - Symbol placeholders: X, “X”, */12, 5/*, X years, X number of layers
   - Any DATE placeholders: DATE, DATE_INSPECTED, DATE_LOSS, DATE_RECEIVED
4. Determine the correct value by interpreting PDF text.
5. Replace ONLY the inner content of parentheses/quotes/symbols (OPTION C).
6. Output ONLY valid JSON mapping:
   {
      "placeholder_string": "filled_value",
      ...
   }
7. NEVER return explanations. ONLY JSON.
"""

    USER_PROMPT = f"""
GLR Template Text:
----------------------
{template_text}

Combined Photo Report Text:
----------------------
{pdf_text}

Return JSON ONLY.
"""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "deepseek/deepseek-r1",
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": USER_PROMPT}
        ],
        "temperature": 0.1
    }

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers=headers,
        data=json.dumps(payload),
    )

    output = response.json()["choices"][0]["message"]["content"]

    try:
        return json.loads(output)
    except:
        match = re.search(r"\{.*\}", output, re.S)
        if match:
            return json.loads(match.group())
        return {}


# ===============================================================
# 4. Fill DOCX template (structure preserved EXACTLY)
# ===============================================================
def fill_template(doc_bytes, mapping):
    doc = docx.Document(doc_bytes)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        for k, v in mapping.items():
            if k in text:
                text = text.replace(k, v)
        paragraph.text = text

    return doc


# ===============================================================
# 5. STREAMLIT UI PIPELINE
# ===============================================================
st.title("🤖 LLM-Powered GLR Auto-Filler (Multi-PDF Support)")

st.write("Upload your **GLR Template (.docx)** and **multiple photo report PDFs**.")

template_file = st.file_uploader("Upload GLR Template", type=["docx"])
photo_files = st.file_uploader("Upload Photo Report PDFs (multiple allowed)", type=["pdf"], accept_multiple_files=True)

if st.button("Generate Filled GLR"):

    if not template_file or not photo_files:
        st.error("Please upload a template and at least one PDF.")
        st.stop()

    # Step 1 — Extract PDF text
    pdf_text = extract_multiple_pdfs(photo_files)

    # Step 2 — Extract template text
    template_doc = docx.Document(template_file)
    template_text = "\n".join([p.text for p in template_doc.paragraphs])

    st.info("🔍 Sending template + PDFs to DeepSeek LLM… Please wait.")

    # Step 3 — LLM builds mapping dictionary
    mapping = call_llm(template_text, pdf_text)

    st.success("🎉 LLM extracted all placeholders!")
    st.json(mapping)

    # Step 4 — Fill template
    filled_doc = fill_template(template_file, mapping)

    # Step 5 — Prepare for download
    buf = BytesIO()
    filled_doc.save(buf)
    buf.seek(0)

    st.download_button(
        "⬇️ Download Completed GLR Report",
        data=buf,
        file_name="Filled_GLR.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
