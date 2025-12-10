import streamlit as st
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from docx.text.run import Run
import requests
import json
import re

st.set_page_config(page_title="Universal GLR Pipeline", layout="wide")
st.title("🔥 Universal Insurance GLR Generator")


# -------------------------------------------------------
# PDF TEXT EXTRACTION
# -------------------------------------------------------
def extract_text_from_pdf(uploaded_pdf):
    text = ""
    try:
        reader = PdfReader(uploaded_pdf)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        text += ""
    return text


# -------------------------------------------------------
# UNIVERSAL PLACEHOLDER DETECTION
# -------------------------------------------------------
def detect_placeholders(doc):
    text = "\n".join([p.text for p in doc.paragraphs])
    patterns = [
        r"\[(.*?)\]",
        r"\$[_]+",
        r"[_]{3,}",
        r"\(\s*Choose an item\.\s*\)",
    ]

    placeholders = set()
    for pat in patterns:
        for m in re.findall(pat, text):
            placeholders.add(m)

    return list(placeholders)


# -------------------------------------------------------
# UNIVERSAL RUN-SAFE TEXT REPLACEMENT
# -------------------------------------------------------
def replace_in_runs(container, search, replace):
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            if search in run.text:
                run.text = run.text.replace(search, replace)

    if hasattr(container, "tables"):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_runs(cell, search, replace)


def fill_template(doc, values):
    # Replace simple [PLACEHOLDER] patterns
    for key, val in values.items():
        search = f"[{key}]"
        replace_in_runs(doc, search, str(val))

    # Replace underscores for narrative expansion
    for key, val in values.items():
        if isinstance(val, str) and len(val) > 20:
            replace_in_runs(doc, "_" * 10, val)

    return doc


# -------------------------------------------------------
# LLM CALL
# -------------------------------------------------------
def call_openrouter(prompt):
    API_KEY = st.secrets["OPENROUTER_API_KEY"]
    url = "https://openrouter.ai/api/v1/chat/completions"

    payload = {
        "model": "",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0
    }

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    response = requests.post(url, json=payload, headers=headers)

    try:
        return response.json()
    except:
        return {"error": "Non-JSON output", "raw": response.text}


# -------------------------------------------------------
# STREAMLIT UI
# -------------------------------------------------------
template_file = st.file_uploader("Upload .docx template", type=["docx"])
uploaded_pdfs = st.file_uploader("Upload Photo Reports", type=["pdf"], accept_multiple_files=True)

if template_file and uploaded_pdfs:

    st.subheader("📄 Extracting PDF text...")
    combined_text = ""

    for pdf in uploaded_pdfs:
        combined_text += extract_text_from_pdf(pdf) + "\n"

    st.text_area("Extracted Text", combined_text[:4000])

    doc = Document(template_file)
    placeholders = detect_placeholders(doc)

    st.subheader("🔍 Detected Placeholders")
    st.write(placeholders)

    if st.button("Run GLR Engine"):

        st.subheader("🤖 Calling LLM")
        prompt = f"""
You are an insurance document automation LLM.

Extract all key–value information from the following report text.
Return clean JSON ONLY with:
- placeholder_name: "value"
- numeric fields in numbers only
- narrative fields in professional complete sentences

Detected placeholders:
{placeholders}

Report text:
{combined_text}
"""

        llm_out = call_openrouter(prompt)
        st.write(llm_out)

        # JSON Extraction
        try:
            text = llm_out.get("choices", [{}])[0].get("message", {}).get("content", "")
            json_str = text[text.find("{"):text.rfind("}") + 1]
            data = json.loads(json_str)
        except Exception as e:
            st.error("JSON Parse Error")
            st.write(e)
            st.stop()

        st.subheader("🧩 Extracted Values")
        st.json(data)

        # Fill Template
        filled_doc = fill_template(doc, data)

        output = BytesIO()
        filled_doc.save(output)
        output.seek(0)

        st.success("✅ Completed GLR Generated!")
        st.download_button(
            label="📥 Download Filled Report",
            data=output,
            file_name="Final_GLR.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Upload template + PDFs to begin.")
